"""
File Text Reader — Extrait le texte brut depuis un fichier stocké (PDF, DOCX, TXT).

Catégorie: file
Mode: sync

Détecte automatiquement le format du fichier à partir de l'extension
ou des magic bytes, puis délègue l'extraction au bon parser.
"""

from __future__ import annotations

import io
from typing import Any

from app.framework.base import BaseTool
from app.framework.schemas import (
    ToolErrorCode,
    ToolExample,
    ToolExecutionMode,
    ToolMetadata,
    ToolParameter,
    ToolResult,
)


class FileTextReader(BaseTool):
    """Extrait le texte depuis un fichier PDF, DOCX ou TXT stocké dans MinIO."""

    @property
    def metadata(self) -> ToolMetadata:
        return ToolMetadata(
            slug="file-text-reader",
            name="File Text Reader",
            description="Extrait le texte brut depuis un fichier stocké (PDF, DOCX, TXT)",
            version="1.0.0",
            category="file",
            execution_mode=ToolExecutionMode.SYNC,
            timeout_seconds=60,
            input_schema=[
                ToolParameter(
                    name="file_key",
                    type="string",
                    required=True,
                    description="Clé de stockage MinIO du fichier à lire",
                ),
            ],
            output_schema=[
                ToolParameter(name="text", type="string", description="Texte extrait du fichier"),
                ToolParameter(name="page_count", type="integer", description="Nombre de pages (PDF uniquement)"),
                ToolParameter(name="format", type="string", description="Format détecté (pdf, docx, txt)"),
            ],
            examples=[
                ToolExample(
                    description="Lire un PDF",
                    input={"file_key": "uploads/contrat.pdf"},
                    output={"text": "Contenu du contrat...", "page_count": 5, "format": "pdf"},
                ),
                ToolExample(
                    description="Lire un DOCX",
                    input={"file_key": "uploads/rapport.docx"},
                    output={"text": "Contenu du rapport...", "format": "docx"},
                ),
            ],
            tags=["file", "reader", "text", "extraction", "pdf", "docx"],
        )

    async def execute(self, params: dict[str, Any], context) -> ToolResult:
        file_key = params.get("file_key", "")
        if not file_key:
            return self.error("file_key requis", ToolErrorCode.INVALID_PARAMS)

        file_data = await context.storage.get(file_key)
        if file_data is None:
            return self.error(f"Fichier introuvable: {file_key}", ToolErrorCode.FILE_NOT_FOUND)

        # Detect format from extension
        key_lower = file_key.lower()
        if key_lower.endswith(".pdf"):
            return await self._read_pdf(file_key, file_data)
        elif key_lower.endswith((".docx", ".doc")):
            return await self._read_docx(file_key, file_data)
        elif key_lower.endswith((".txt", ".md", ".csv")):
            return self._read_text(file_key, file_data)

        # Fallback: try magic bytes
        if file_data[:5] == b"%PDF-":
            return await self._read_pdf(file_key, file_data)
        if file_data[:4] == b"PK\x03\x04":
            return await self._read_docx(file_key, file_data)

        # Last resort: try as plain text
        return self._read_text(file_key, file_data)

    async def _read_pdf(self, file_key: str, file_data: bytes) -> ToolResult:
        from PyPDF2 import PdfReader

        try:
            reader = PdfReader(io.BytesIO(file_data))
        except Exception as e:
            return self.error(f"Fichier PDF invalide: {e}", ToolErrorCode.PROCESSING_ERROR)

        pages_text = []
        for page in reader.pages:
            text = page.extract_text() or ""
            pages_text.append(text)

        return self.success({
            "text": "\n\n".join(pages_text),
            "page_count": len(reader.pages),
            "format": "pdf",
        })

    async def _read_docx(self, file_key: str, file_data: bytes) -> ToolResult:
        from docx import Document

        try:
            doc = Document(io.BytesIO(file_data))
        except Exception as e:
            return self.error(f"Fichier DOCX invalide: {e}", ToolErrorCode.PROCESSING_ERROR)

        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)

        return self.success({
            "text": "\n".join(parts),
            "format": "docx",
        })

    def _read_text(self, file_key: str, file_data: bytes) -> ToolResult:
        try:
            text = file_data.decode("utf-8")
        except UnicodeDecodeError:
            try:
                text = file_data.decode("latin-1")
            except Exception as e:
                return self.error(f"Impossible de décoder le fichier: {e}", ToolErrorCode.PROCESSING_ERROR)

        return self.success({
            "text": text,
            "format": "txt",
        })
