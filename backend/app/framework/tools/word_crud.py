"""
Word CRUD — Créer, lire, modifier et supprimer des fichiers DOCX.

Catégorie: file
Mode: sync

Actions:
    create  → Génère un DOCX depuis des données structurées (titre, paragraphes, tableaux)
    read    → Extrait le contenu d'un DOCX (texte, paragraphes, tableaux)
    update  → Modifie un DOCX (ajouter du contenu, remplacer du texte)
    delete  → Supprime un fichier DOCX du stockage
"""

from __future__ import annotations

import io
from typing import Any

from app.framework.base import BaseTool
from app.framework.schemas import (
    ToolErrorCode,
    ToolExample,
    ToolExecutionMode,
    ToolMetadata,
    ToolParameter,
    ToolResult,
)


class WordCrud(BaseTool):
    """CRUD complet pour les fichiers DOCX via MinIO."""

    @property
    def metadata(self) -> ToolMetadata:
        return ToolMetadata(
            slug="word-crud",
            name="Word CRUD",
            description="Créer, lire, modifier et supprimer des fichiers Word (DOCX)",
            version="1.0.0",
            category="file",
            execution_mode=ToolExecutionMode.SYNC,
            timeout_seconds=30,
            input_schema=[
                ToolParameter(name="action", type="string", required=True,
                              description="Action: create, read, update, delete"),
                ToolParameter(name="storage_key", type="string",
                              description="Chemin MinIO du fichier"),
                ToolParameter(name="data", type="object",
                              description="Contenu: {title, paragraphs: [{text, style}], tables: [{headers, rows}]}"),
                ToolParameter(name="options", type="object",
                              description="Options: {font_name, font_size}"),
            ],
            output_schema=[
                ToolParameter(name="storage_key", type="string"),
                ToolParameter(name="text", type="string", description="Texte complet extrait"),
                ToolParameter(name="paragraphs", type="array"),
                ToolParameter(name="tables", type="array"),
                ToolParameter(name="paragraph_count", type="integer"),
                ToolParameter(name="table_count", type="integer"),
            ],
            examples=[
                ToolExample(
                    description="Créer un document Word",
                    input={
                        "action": "create",
                        "storage_key": "docs/rapport.docx",
                        "data": {
                            "title": "Rapport Mensuel",
                            "paragraphs": [
                                {"text": "Introduction du rapport.", "style": "Normal"},
                                {"text": "Section 1: Résultats", "style": "Heading 1"},
                                {"text": "Les résultats sont positifs.", "style": "Normal"},
                            ],
                        },
                    },
                    output={"storage_key": "docs/rapport.docx", "paragraph_count": 4},
                ),
                ToolExample(
                    description="Lire un document Word",
                    input={"action": "read", "storage_key": "docs/rapport.docx"},
                    output={
                        "text": "Rapport Mensuel\nIntroduction...",
                        "paragraph_count": 4,
                        "table_count": 0,
                    },
                ),
            ],
            tags=["file", "word", "docx", "crud", "office"],
        )

    async def execute(self, params: dict[str, Any], context) -> ToolResult:
        action = params.get("action", "")
        if action == "create":
            return await self._create(params, context)
        elif action == "read":
            return await self._read(params, context)
        elif action == "update":
            return await self._update(params, context)
        elif action == "delete":
            return await self._delete(params, context)
        else:
            return self.error(
                f"Action inconnue: '{action}'. Actions: create, read, update, delete",
                ToolErrorCode.INVALID_PARAMS,
            )

    async def _create(self, params: dict[str, Any], context) -> ToolResult:
        from docx import Document
        from docx.shared import Pt

        storage_key = params.get("storage_key", "")
        data = params.get("data", {})
        options = params.get("options", {})

        if not storage_key:
            return self.error("storage_key requis pour create", ToolErrorCode.INVALID_PARAMS)

        doc = Document()

        # Style par défaut
        font_name = options.get("font_name", "Calibri")
        font_size = options.get("font_size", 11)
        style = doc.styles["Normal"]
        style.font.name = font_name
        style.font.size = Pt(font_size)

        # Titre
        title = data.get("title")
        if title:
            doc.add_heading(title, level=0)

        # Paragraphes
        paragraphs = data.get("paragraphs", [])
        for para in paragraphs:
            if isinstance(para, str):
                doc.add_paragraph(para)
            elif isinstance(para, dict):
                text = para.get("text", "")
                para_style = para.get("style", "Normal")
                if para_style.startswith("Heading"):
                    level = int(para_style.split()[-1]) if para_style != "Heading" else 1
                    doc.add_heading(text, level=level)
                else:
                    doc.add_paragraph(text, style=para_style)

        # Tableaux
        tables = data.get("tables", [])
        for table_data in tables:
            headers = table_data.get("headers", [])
            rows = table_data.get("rows", [])
            if headers:
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = "Table Grid"
                for i, header in enumerate(headers):
                    table.rows[0].cells[i].text = str(header)
                for row_data in rows:
                    row = table.add_row()
                    for i, cell_value in enumerate(row_data):
                        if i < len(headers):
                            row.cells[i].text = str(cell_value)

        # Sauvegarder
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        await context.storage.put(
            storage_key, buffer.getvalue(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        return self.success({
            "storage_key": storage_key,
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
        })

    async def _read(self, params: dict[str, Any], context) -> ToolResult:
        from docx import Document

        storage_key = params.get("storage_key", "")
        if not storage_key:
            return self.error("storage_key requis pour read", ToolErrorCode.INVALID_PARAMS)

        file_data = await context.storage.get(storage_key)
        if file_data is None:
            return self.error(f"Fichier introuvable: {storage_key}", ToolErrorCode.FILE_NOT_FOUND)

        try:
            doc = Document(io.BytesIO(file_data))
        except Exception as e:
            return self.error(f"Fichier DOCX invalide: {e}", ToolErrorCode.PROCESSING_ERROR)

        # Extraire paragraphes
        paragraphs = []
        full_text_parts = []
        for para in doc.paragraphs:
            paragraphs.append({
                "text": para.text,
                "style": para.style.name if para.style else "Normal",
            })
            if para.text.strip():
                full_text_parts.append(para.text)

        # Extraire tableaux
        tables = []
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                table_rows.append([cell.text for cell in row.cells])
            headers = table_rows[0] if table_rows else []
            tables.append({
                "headers": headers,
                "rows": table_rows[1:] if len(table_rows) > 1 else [],
            })

        return self.success({
            "storage_key": storage_key,
            "text": "\n".join(full_text_parts),
            "paragraphs": paragraphs,
            "tables": tables,
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
        })

    async def _update(self, params: dict[str, Any], context) -> ToolResult:
        from docx import Document

        storage_key = params.get("storage_key", "")
        data = params.get("data", {})

        if not storage_key:
            return self.error("storage_key requis pour update", ToolErrorCode.INVALID_PARAMS)

        file_data = await context.storage.get(storage_key)
        if file_data is None:
            return self.error(f"Fichier introuvable: {storage_key}", ToolErrorCode.FILE_NOT_FOUND)

        try:
            doc = Document(io.BytesIO(file_data))
        except Exception as e:
            return self.error(f"Fichier DOCX invalide: {e}", ToolErrorCode.PROCESSING_ERROR)

        # Remplacer du texte
        replacements = data.get("replace", {})
        for old_text, new_text in replacements.items():
            for para in doc.paragraphs:
                if old_text in para.text:
                    for run in para.runs:
                        if old_text in run.text:
                            run.text = run.text.replace(old_text, new_text)

        # Ajouter des paragraphes
        add_paragraphs = data.get("add_paragraphs", [])
        for para in add_paragraphs:
            if isinstance(para, str):
                doc.add_paragraph(para)
            elif isinstance(para, dict):
                text = para.get("text", "")
                style = para.get("style", "Normal")
                if style.startswith("Heading"):
                    level = int(style.split()[-1]) if style != "Heading" else 1
                    doc.add_heading(text, level=level)
                else:
                    doc.add_paragraph(text, style=style)

        # Ajouter des tableaux
        add_tables = data.get("add_tables", [])
        for table_data in add_tables:
            headers = table_data.get("headers", [])
            rows = table_data.get("rows", [])
            if headers:
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = "Table Grid"
                for i, header in enumerate(headers):
                    table.rows[0].cells[i].text = str(header)
                for row_data in rows:
                    row = table.add_row()
                    for i, cell_value in enumerate(row_data):
                        if i < len(headers):
                            row.cells[i].text = str(cell_value)

        # Sauvegarder
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        await context.storage.put(
            storage_key, buffer.getvalue(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        return self.success({
            "storage_key": storage_key,
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
        })

    async def _delete(self, params: dict[str, Any], context) -> ToolResult:
        storage_key = params.get("storage_key", "")
        if not storage_key:
            return self.error("storage_key requis pour delete", ToolErrorCode.INVALID_PARAMS)

        deleted = await context.storage.delete(storage_key)
        if not deleted:
            return self.error(f"Fichier introuvable: {storage_key}", ToolErrorCode.FILE_NOT_FOUND)

        return self.success({"storage_key": storage_key, "deleted": True})
