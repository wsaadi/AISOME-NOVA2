"""
Document Generator — Génère des documents (DOCX, PDF) depuis du contenu structuré.

Catégorie: file
Mode: sync

Reçoit un contenu structuré (titre, sections, format) et produit
un fichier document stocké dans MinIO.
"""

from __future__ import annotations

import io
from typing import Any

from app.framework.base import BaseTool
from app.framework.schemas import (
    ToolErrorCode,
    ToolExample,
    ToolExecutionMode,
    ToolMetadata,
    ToolParameter,
    ToolResult,
)


class DocumentGenerator(BaseTool):
    """Génère des documents DOCX ou PDF depuis du contenu structuré."""

    @property
    def metadata(self) -> ToolMetadata:
        return ToolMetadata(
            slug="document-generator",
            name="Document Generator",
            description="Génère des documents DOCX ou PDF depuis du contenu structuré",
            version="1.0.0",
            category="file",
            execution_mode=ToolExecutionMode.SYNC,
            timeout_seconds=30,
            input_schema=[
                ToolParameter(
                    name="format",
                    type="string",
                    required=True,
                    description="Format de sortie: docx ou pdf",
                ),
                ToolParameter(
                    name="content",
                    type="object",
                    required=True,
                    description="Contenu structuré: {title, sections: [{title, content}], format: {font, title_size, heading_size, body_size, margins}}",
                ),
                ToolParameter(
                    name="output_key",
                    type="string",
                    required=True,
                    description="Clé de stockage MinIO pour le fichier généré",
                ),
            ],
            output_schema=[
                ToolParameter(name="file_key", type="string", description="Clé de stockage du fichier généré"),
                ToolParameter(name="format", type="string", description="Format du document généré"),
            ],
            examples=[
                ToolExample(
                    description="Générer un rapport Word",
                    input={
                        "format": "docx",
                        "content": {
                            "title": "Rapport d'analyse",
                            "sections": [
                                {"title": "Introduction", "content": "Texte d'intro..."},
                                {"title": "Résultats", "content": "Texte des résultats..."},
                            ],
                        },
                        "output_key": "reports/rapport.docx",
                    },
                    output={"file_key": "reports/rapport.docx", "format": "docx"},
                ),
            ],
            tags=["file", "generator", "document", "docx", "pdf", "report"],
        )

    async def execute(self, params: dict[str, Any], context) -> ToolResult:
        fmt = params.get("format", "")
        content = params.get("content", {})
        output_key = params.get("output_key", "")

        if not fmt:
            return self.error("format requis (docx ou pdf)", ToolErrorCode.INVALID_PARAMS)
        if not content:
            return self.error("content requis", ToolErrorCode.INVALID_PARAMS)
        if not output_key:
            return self.error("output_key requis", ToolErrorCode.INVALID_PARAMS)

        if fmt == "docx":
            return await self._generate_docx(content, output_key, context)
        elif fmt == "pdf":
            return await self._generate_pdf(content, output_key, context)
        else:
            return self.error(f"Format non supporté: '{fmt}'. Formats: docx, pdf", ToolErrorCode.INVALID_PARAMS)

    async def _generate_docx(self, content: dict, output_key: str, context) -> ToolResult:
        from docx import Document
        from docx.shared import Pt, Cm

        doc = Document()

        # Apply formatting options
        fmt = content.get("format", {})
        font_name = fmt.get("font", "Calibri")
        body_size = fmt.get("body_size", 11)

        style = doc.styles["Normal"]
        style.font.name = font_name
        style.font.size = Pt(body_size)

        # Set margins
        margins = fmt.get("margins", {})
        for section in doc.sections:
            section.top_margin = Cm(margins.get("top", 2.5))
            section.bottom_margin = Cm(margins.get("bottom", 2.5))
            section.left_margin = Cm(margins.get("left", 2.5))
            section.right_margin = Cm(margins.get("right", 2.5))

        # Title
        title = content.get("title")
        if title:
            doc.add_heading(title, level=0)

        # Sections
        sections = content.get("sections", [])
        for section_data in sections:
            section_title = section_data.get("title")
            section_content = section_data.get("content", "")

            if section_title:
                doc.add_heading(section_title, level=1)

            # Split content by newlines and add as paragraphs
            for paragraph in section_content.split("\n"):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        await context.storage.put(
            output_key,
            buffer.getvalue(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        return self.success({
            "file_key": output_key,
            "format": "docx",
        })

    async def _generate_pdf(self, content: dict, output_key: str, context) -> ToolResult:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        fmt = content.get("format", {})
        body_size = fmt.get("body_size", 11)
        margins = fmt.get("margins", {})
        margin_val = margins.get("left", 2.5)

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            leftMargin=margin_val * cm,
            rightMargin=margins.get("right", margin_val) * cm,
            topMargin=margins.get("top", margin_val) * cm,
            bottomMargin=margins.get("bottom", margin_val) * cm,
        )

        styles = getSampleStyleSheet()
        title_style = styles["Title"]
        heading_style = ParagraphStyle(
            "GenHeading",
            parent=styles["Heading1"],
            fontSize=fmt.get("heading_size", 14),
            spaceAfter=12,
            spaceBefore=12,
        )
        body_style = ParagraphStyle(
            "GenBody",
            parent=styles["Normal"],
            fontSize=body_size,
            leading=body_size * 1.4,
            spaceAfter=6,
        )

        elements = []

        # Title
        title = content.get("title")
        if title:
            elements.append(Paragraph(title, title_style))
            elements.append(Spacer(1, 0.5 * cm))

        # Sections
        sections = content.get("sections", [])
        for section_data in sections:
            section_title = section_data.get("title")
            section_content = section_data.get("content", "")

            if section_title:
                elements.append(Paragraph(section_title, heading_style))

            for paragraph in section_content.split("\n"):
                if paragraph.strip():
                    elements.append(Paragraph(paragraph.strip(), body_style))

        if not elements:
            elements.append(Paragraph("&nbsp;", body_style))

        doc.build(elements)
        buffer.seek(0)

        await context.storage.put(output_key, buffer.getvalue(), "application/pdf")

        return self.success({
            "file_key": output_key,
            "format": "pdf",
        })
